# -*- coding: utf-8 -*-
"""Fix database table descriptions in the docx file to match actual DB schema."""

from docx import Document
import re

INPUT = r"C:\Users\zzy\Desktop\学校\课程\智能软件开发综合实训\会议纪要智能体\docx_output\课程报告-最终版v8.docx"
OUTPUT = r"C:\Users\zzy\Desktop\学校\课程\智能软件开发综合实训\会议纪要智能体\docx_output\课程报告-最终版v9.docx"

doc = Document(INPUT)

# === Fix 1: Update field counts in text ===
# meetings: 11 -> 13, meeting_minutes: 10 -> 11
replacements = [
    ("meetings（会议基本信息，11个字段）", "meetings（会议基本信息，13个字段）"),
    ("meeting_minutes（纪要内容，10个字段）", "meeting_minutes（纪要内容，11个字段）"),
]

def replace_in_paragraph(para, old, new):
    """Replace text across runs preserving formatting."""
    full_text = "".join(r.text for r in para.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    # Clear all runs, set first run to new text
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    return True

def replace_in_table(table, old, new):
    """Replace text in table cells."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_in_paragraph(para, old, new)

# Apply text replacements
for old, new in replacements:
    for para in doc.paragraphs:
        replace_in_paragraph(para, old, new)
    for table in doc.tables:
        replace_in_table(table, old, new)

# === Fix 2: Fix the Meeting model code snippet ===
# The code snippet in section 3.3 is missing duration, input_type, updated_at
# We need to find the paragraph containing the Meeting model code and update it

OLD_MEETING_CODE = """class Meeting(Base):
    \"\"\"会议记录表\"\"\"
    __tablename__ = "meetings"
    id = Column(Integer, primary_key=True, autoincrement=True)
    title = Column(String(200), nullable=False)
    meeting_date = Column(DateTime, default=datetime.now)
    participants = Column(String(500))
    raw_text = Column(Text)
    audio_path = Column(String(500))
    status = Column(String(20), default="processing")
    is_public = Column(Boolean, default=False)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.now)

    creator = relationship("User")
    minutes = relationship("MeetingMinutes", back_populates="meeting",
                           uselist=False, cascade="all, delete-orphan")
    action_items = relationship("ActionItem", back_populates="meeting",
                                cascade="all, delete-orphan")"""

NEW_MEETING_CODE = """class Meeting(Base):
    \"\"\"会议记录表\"\"\"
    __tablename__ = "meetings"
    id = Column(Integer, primary_key=True, autoincrement=True)
    title = Column(String(200), nullable=False)
    meeting_date = Column(DateTime, default=datetime.now)
    duration = Column(String(50))
    participants = Column(String(500))
    input_type = Column(String(20), default="text")
    raw_text = Column(Text)
    audio_path = Column(String(500))
    status = Column(String(20), default="processing")
    is_public = Column(Boolean, default=False)
    created_by = Column(Integer, ForeignKey("users.id"))
    created_at = Column(DateTime, default=datetime.now)
    updated_at = Column(DateTime, default=datetime.now,
                        onupdate=datetime.now)

    creator = relationship("User")
    minutes = relationship("MeetingMinutes", back_populates="meeting",
                           uselist=False, cascade="all, delete-orphan")
    action_items = relationship("ActionItem", back_populates="meeting",
                                cascade="all, delete-orphan")"""

# Search for the code block containing the Meeting model
found = False
for i, para in enumerate(doc.paragraphs):
    if "class Meeting(Base):" in para.text and '__tablename__ = "meetings"' in para.text:
        # This paragraph contains the Meeting model code
        # We need to check if it's a code-style paragraph
        full_text = para.text
        if "duration" not in full_text:
            # Need to replace
            new_text = full_text.replace(
                '    meeting_date = Column(DateTime, default=datetime.now)\n    participants = Column(String(500))',
                '    meeting_date = Column(DateTime, default=datetime.now)\n    duration = Column(String(50))\n    participants = Column(String(500))\n    input_type = Column(String(20), default="text")'
            )
            new_text = new_text.replace(
                '    created_at = Column(DateTime, default=datetime.now)\n\n    creator',
                '    created_at = Column(DateTime, default=datetime.now)\n    updated_at = Column(DateTime, default=datetime.now,\n                        onupdate=datetime.now)\n\n    creator'
            )
            if para.runs:
                para.runs[0].text = new_text
                for r in para.runs[1:]:
                    r.text = ""
                found = True
                print(f"Fixed Meeting model code at paragraph {i}")
            break

if not found:
    print("WARNING: Could not find Meeting model code snippet to fix")
    # Try a more flexible search
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if "Meeting(Base)" in text and "participants" in text and "duration" not in text:
            print(f"  Found candidate paragraph {i}: {text[:100]}...")

# === Fix 3: Fix the text description of Meeting model fields ===
# "Meeting会议表包含11个字段" -> "13个字段"
field_count_replacements = [
    ("Meeting会议表包含11个字段", "Meeting会议表包含13个字段"),
    ("MeetingMinutes纪要表包含10个字段", "MeetingMinutes纪要表包含11个字段"),
]

for old, new in field_count_replacements:
    for para in doc.paragraphs:
        replace_in_paragraph(para, old, new)
    for table in doc.tables:
        replace_in_table(table, old, new)

doc.save(OUTPUT)
print(f"Saved to {OUTPUT}")

# Verify
doc2 = Document(OUTPUT)
for para in doc2.paragraphs:
    if "13个字段" in para.text or "11个字段" in para.text:
        print(f"  Verified: {para.text[:100]}")
    if "class Meeting" in para.text and "duration" in para.text:
        print(f"  Verified Meeting model has duration field")
