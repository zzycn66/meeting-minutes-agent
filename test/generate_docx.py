# -*- coding: utf-8 -*-
"""生成代码讲解 Word 文档的脚本"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_page(doc):
    """设置页面为 A4 纸张"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)

def setup_styles(doc):
    """设置文档样式"""
    # 正文样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 代码样式
    if "Code" not in [s.name for s in doc.styles]:
        code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(3)
        code_style.paragraph_format.space_after = Pt(3)
        code_style.paragraph_format.left_indent = Cm(0.5)
    
    # 标题样式
    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "微软雅黑"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def add_code_block(doc, code_text, language="python"):
    """添加代码块"""
    # 添加代码标题
    p = doc.add_paragraph()
    run = p.add_run(f"代码块 ({language})：")
    run.bold = True
    run.font.size = Pt(10)
    
    # 添加代码内容
    for line in code_text.split("\n"):
        p = doc.add_paragraph(line, style="Code")
        # 设置中文字体
        for run in p.runs:
            run.font.name = "Consolas"
            # 设置中文备用字体
            from docx.oxml.ns import qn
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def add_explanation(doc, text):
    """添加解释说明"""
    p = doc.add_paragraph()
    run = p.add_run("讲解：")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    p.add_run(text)

def add_bullet(doc, text, level=0):
    """添加要点"""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.0 * level)

def generate_docx():
    """生成 Word 文档"""
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    
    # 封面
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph("会议纪要智能体", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph("代码详细讲解文档", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    info = doc.add_paragraph("智能软件开发综合实训课程")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # 目录页
    doc.add_heading("目录", level=1)
    doc.add_paragraph("第一章：项目入口 main.py")
    doc.add_paragraph("第二章：配置系统 config.py")
    doc.add_paragraph("第三章：数据库层 database.py + models.py")
    doc.add_paragraph("第四章：语音识别模块 speech_service.py")
    doc.add_paragraph("第五章：NLP信息提取 nlp_service.py")
    doc.add_paragraph("第六章：智能问答模块 qa_service.py")
    doc.add_paragraph("第七章：音频增强模块 audio_processor.py")
    doc.add_paragraph("第八章：说话人识别 speaker_diarization.py")
    doc.add_paragraph("第九章：实时转写模块 streaming_service.py")
    doc.add_paragraph("第十章：API路由层 routes/*.py")
    doc.add_paragraph("第十一章：前端交互逻辑 app.js")
    
    doc.add_page_break()
    
    # ==================== 第一章 ====================
    doc.add_heading("第一章：项目入口 main.py", level=1)
    
    doc.add_heading("1.1 环境预配置", level=2)
    
    add_code_block(doc, '''# -*- coding: utf-8 -*-
"""会议纪要智能体 - 主应用入口"""
import os
import sys
import logging''')
    
    add_explanation(doc, "第一行声明文件编码为 UTF-8，确保中文注释和字符串能正确处理。导入标准库：os 用于文件操作，sys 用于系统操作，logging 用于日志。")
    
    doc.add_heading("1.2 Python 路径配置", level=2)
    
    add_code_block(doc, '''_backend_dir = os.path.dirname(os.path.abspath(__file__))
if _backend_dir not in sys.path:
    sys.path.insert(0, _backend_dir)''')
    
    add_explanation(doc, """
获取当前文件所在目录的绝对路径，并将项目根目录添加到 sys.path。
这样做的目的是确保能正确导入 app 包，支持从任意位置运行脚本。
- os.path.abspath(__file__)：获取当前文件的绝对路径
- os.path.dirname(...)：获取目录部分
- sys.path.insert(0, ...)：将路径插入到最前面，优先查找
""")
    
    doc.add_heading("1.3 环境变量加载", level=2)
    
    add_code_block(doc, '''# 提前加载 .env 并设置 HF_ENDPOINT
_dotenv_path = os.path.join(_backend_dir, ".env")
if os.path.exists(_dotenv_path):
    try:
        from dotenv import load_dotenv
        load_dotenv(_dotenv_path)
    except ImportError:
        pass''')
    
    add_explanation(doc, """
从 .env 文件加载环境变量。.env 文件存储敏感配置（如 API Key），不上传到 Git。
关键点：必须在导入 transformers 之前加载，因为 HuggingFace 库在导入时就缓存了环境变量。
""")
    
    doc.add_heading("1.4 HuggingFace 镜像设置", level=2)
    
    add_code_block(doc, '''if "HF_ENDPOINT" not in os.environ:
    os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"''')
    
    add_explanation(doc, """
设置 HuggingFace 镜像地址为国内镜像 hf-mirror.com。
原因：国内直接访问 huggingface.co 很慢或无法访问，使用镜像站加速模型下载。
""")
    
    doc.add_heading("1.5 深度学习框架预加载", level=2)
    
    add_code_block(doc, '''try:
    from transformers import AutoTokenizer, AutoModelForCausalLM
    import torch
    logging.getLogger(__name__).info("transformers/torch 预加载成功")
except ImportError as e:
    logging.getLogger(__name__).warning(f"transformers/torch 预加载失败: {e}")''')
    
    add_explanation(doc, """
提前导入深度学习框架，避免首次请求时的延迟。
- AutoTokenizer：自动加载分词器
- AutoModelForCausalLM：自动加载因果语言模型（如 Qwen3）
- torch：PyTorch 深度学习框架
如果导入失败，只记录警告，不阻塞服务启动。
""")
    
    doc.add_heading("1.6 FastAPI 应用创建", level=2)
    
    add_code_block(doc, '''from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(
    title="会议纪要智能体",
    description="智能会议纪要助手 - 自动转写、提取、生成结构化会议纪要",
    version="1.0.0",
)''')
    
    add_explanation(doc, """
创建 FastAPI 应用实例。
- FastAPI：高性能异步 Web 框架
- title/description/version：会显示在 API 文档页面（访问 /docs 可以看到）
- 自动生成 Swagger/OpenAPI 文档
""")
    
    doc.add_heading("1.7 CORS 跨域配置", level=2)
    
    add_code_block(doc, '''app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)''')
    
    add_explanation(doc, """
CORS（跨域资源共享）：浏览器安全策略，前端默认不能访问不同域名的后端 API。
- allow_origins：允许哪些域名访问
- allow_credentials：允许携带 Cookie（用户认证需要）
- allow_methods：允许所有 HTTP 方法（GET/POST/PUT/DELETE）
- allow_headers：允许所有请求头
""")
    
    doc.add_heading("1.8 路由注册", level=2)
    
    add_code_block(doc, '''app.include_router(meetings.router)    # 会议管理 API
app.include_router(minutes.router)     # 纪要管理 API
app.include_router(export.router)      # 导出 API
app.include_router(auth.router)        # 认证 API
app.include_router(templates.router)   # 模板 API
app.include_router(websocket.router)   # WebSocket API
app.include_router(audio.audio_router) # 音频处理 API
app.include_router(public.router)      # 公开文档 API
app.include_router(qa_router)          # 智能问答 API''')
    
    add_explanation(doc, """
include_router：将路由模块注册到主应用。
每个路由模块定义了一组相关的 API 端点。
例如 meetings.router 处理 /api/meetings/* 的请求。
""")
    
    doc.add_heading("1.9 启动时初始化", level=2)
    
    add_code_block(doc, '''@app.on_event("startup")
async def startup():
    Base.metadata.create_all(bind=engine)
    
    import threading
    def _preload():
        from app.services.speech_service import _get_faster_whisper_model
        m = _get_faster_whisper_model()
    threading.Thread(target=_preload, daemon=True).start()''')
    
    add_explanation(doc, """
@app.on_event("startup")：FastAPI 生命周期钩子，服务启动时自动执行。
Base.metadata.create_all：根据模型定义创建数据库表（如果不存在）。
模型预加载：使用后台线程加载 Whisper 模型，避免阻塞服务启动。
- daemon=True：主线程退出时自动终止子线程
- Whisper 模型加载需要几秒到几十秒，这样服务启动后就能立即接受请求
""")
    
    doc.add_heading("1.10 首页路由", level=2)
    
    add_code_block(doc, '''@app.get("/")
async def index():
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path, headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        })
    return {"message": "会议纪要智能体API服务运行中", "docs": "/docs"}''')
    
    add_explanation(doc, """
返回前端首页 index.html。
缓存控制头：确保用户总是获取最新版本的前端代码。
- Cache-Control: no-cache：每次请求都要验证
- Pragma: no-cache：HTTP/1.0 兼容
- Expires: 0：立即过期
""")
    
    doc.add_page_break()
    
    # ==================== 第二章 ====================
    doc.add_heading("第二章：配置系统 config.py", level=1)
    
    doc.add_heading("2.1 环境变量加载", level=2)
    
    add_code_block(doc, '''import os
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env"))''')
    
    add_explanation(doc, """
加载 .env 文件中的环境变量。
.env 文件示例：
DEEPSEEK_API_KEY=sk-xxxx
MIMO_API_KEY=xxxx
WHISPER_MODEL_SIZE=medium
""")
    
    doc.add_heading("2.2 数据库配置", level=2)
    
    add_code_block(doc, '''BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATABASE_URL = f"sqlite:///{os.path.join(BASE_DIR, 'meeting_agent.db')}"''')
    
    add_explanation(doc, """
BASE_DIR：项目根目录的绝对路径。
DATABASE_URL：SQLite 数据库连接字符串。
格式：sqlite:///数据库文件路径
SQLite 优势：轻量级、免安装、适合中小型应用。
""")
    
    doc.add_heading("2.3 语音识别配置", level=2)
    
    add_code_block(doc, '''# MiMo ASR 配置
MIMO_API_KEY = os.environ.get("MIMO_API_KEY", "")
MIMO_ASR_BASE_URL = os.environ.get("MIMO_ASR_BASE_URL", "https://api.xiaomimimo.com/v1")

# Whisper 模型配置
WHISPER_MODEL_SIZE = os.environ.get("WHISPER_MODEL_SIZE", "medium")''')
    
    add_explanation(doc, """
MiMo ASR：小米的云端语音识别服务。
os.environ.get("key", "default")：优先读取环境变量，不存在则使用默认值。
Whisper 模型大小选择：
- tiny: 39M, ~1GB 内存, 最快, 准确率最低
- base: 74M, ~1GB 内存, 快, 准确率低
- small: 244M, ~2GB 内存, 中等, 准确率中等
- medium: 769M, ~5GB 内存, 较慢, 准确率高（推荐）
- large-v3: 1550M, ~10GB 内存, 最慢, 准确率最高
""")
    
    doc.add_heading("2.4 大语言模型配置", level=2)
    
    add_code_block(doc, '''# DeepSeek API 配置
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")

# Qwen3 本地模型配置
LOCAL_MODEL_NAME = os.environ.get("LOCAL_MODEL_NAME", "Qwen/Qwen3-0.6B")
LOCAL_MODEL_DEVICE = os.environ.get("LOCAL_MODEL_DEVICE", "cpu")
LOCAL_MODEL_MAX_TOKENS = int(os.environ.get("LOCAL_MODEL_MAX_TOKENS", "1024"))''')
    
    add_explanation(doc, """
DeepSeek：大语言模型 API 服务，API 兼容 OpenAI 格式。
Qwen3：通义千问 3 代本地模型。
- 0.6B: 6亿参数, ~2GB 内存, 快, 准确率一般
- 1.7B: 17亿参数, ~4GB 内存, 中等
- 4B: 40亿参数, ~8GB 内存, 较慢
LOCAL_MODEL_DEVICE：运行设备
- cpu：CPU 推理（较慢，但无需 GPU）
- cuda：GPU 推理（需要 NVIDIA GPU）
""")
    
    doc.add_page_break()
    
    # ==================== 第三章 ====================
    doc.add_heading("第三章：数据库层", level=1)
    
    doc.add_heading("3.1 数据库连接 database.py", level=2)
    
    add_code_block(doc, '''engine = create_engine(
    DATABASE_URL,
    connect_args={"check_same_thread": False},
    pool_pre_ping=True,
    pool_size=5,
    max_overflow=5,
    pool_recycle=3600,
    echo=False,
)''')
    
    add_explanation(doc, """
参数详解：
- check_same_thread=False：SQLite 默认不允许跨线程访问，此参数禁用此检查
- pool_pre_ping=True：每次使用连接前检测连接是否有效
- pool_size=5：连接池保持 5 个空闲连接
- max_overflow=5：超出 pool_size 的最大连接数（最多 10 个并发连接）
- pool_recycle=3600：连接使用超过 1 小时后回收
- echo=False：不打印 SQL 语句（生产环境）
""")
    
    doc.add_heading("3.2 SQLite 性能优化", level=2)
    
    add_code_block(doc, '''@event.listens_for(engine, "connect")
def _set_sqlite_pragma(dbapi_connection, connection_record):
    cursor = dbapi_connection.cursor()
    cursor.execute("PRAGMA journal_mode=WAL")
    cursor.execute("PRAGMA synchronous=NORMAL")
    cursor.execute("PRAGMA cache_size=-64000")
    cursor.execute("PRAGMA temp_store=MEMORY")
    cursor.close()''')
    
    add_explanation(doc, """
PRAGMA 是 SQLite 的配置命令：
- journal_mode=WAL：启用 WAL 模式，允许读写并发，性能提升 2-5 倍
- synchronous=NORMAL：平衡性能和安全
- cache_size=-64000：64MB 缓存（负数表示 KB）
- temp_store=MEMORY：临时表存储在内存（更快）

WAL 模式优势：
- 普通模式：读写互斥，写操作阻塞读操作
- WAL 模式：读写可以并发，性能大幅提升
""")
    
    doc.add_heading("3.3 数据库会话管理", level=2)
    
    add_code_block(doc, '''SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()''')
    
    add_explanation(doc, """
- sessionmaker：创建会话类
  - autocommit=False：不自动提交，需要手动 db.commit()
  - autoflush=False：不自动刷新
- declarative_base()：声明式基类，所有模型继承此类
- get_db：FastAPI 依赖注入函数
  - yield：返回数据库会话，同时保持函数运行
  - finally：确保会话关闭（无论成功或失败）
""")
    
    doc.add_page_break()
    
    # ==================== 第四章 ====================
    doc.add_heading("第四章：语音识别模块 speech_service.py", level=1)
    
    doc.add_heading("4.1 Whisper 模型加载", level=2)
    
    add_code_block(doc, '''_faster_whisper_model = None

def _get_faster_whisper_model():
    global _faster_whisper_model
    if _faster_whisper_model is not None:
        return _faster_whisper_model

    try:
        from faster_whisper import WhisperModel
        compute_type = "int8" if os.environ.get("LOCAL_MODEL_DEVICE", "cpu") == "cpu" else "float16"
        _faster_whisper_model = WhisperModel(
            WHISPER_MODEL_SIZE,
            device=os.environ.get("LOCAL_MODEL_DEVICE", "cpu"),
            compute_type=compute_type,
        )
        return _faster_whisper_model
    except ImportError:
        return None''')
    
    add_explanation(doc, """
单例模式：使用全局变量缓存模型，避免重复加载。
faster_whisper：基于 CTranslate2 加速的 Whisper，速度提升 2-4 倍。
- compute_type="int8"：CPU 使用 int8 量化，更快
- compute_type="float16"：GPU 使用 float16，更准
""")
    
    doc.add_heading("4.2 Whisper 识别核心", level=2)
    
    add_code_block(doc, '''def _recognize_with_whisper(self, filepath: str) -> dict:
    model = _get_faster_whisper_model()
    
    if model:
        segments, info = model.transcribe(
            filepath,
            language="zh",
            beam_size=5,
            best_of=5,
            temperature=0.0,
            condition_on_previous_text=False,
            initial_prompt=WHISPER_INITIAL_PROMPT,
        )
        
        text = "".join([seg.text for seg in segments])
        return {"success": True, "text": text}''')
    
    add_explanation(doc, """
参数说明：
- language="zh"：指定中文，提升识别准确率
- beam_size=5：束搜索宽度，越大越准但越慢
- best_of=5：生成 5 个候选，选择最佳
- temperature=0.0：贪心解码，确定性输出
- condition_on_previous_text=False：不依赖前文，防止幻觉累积
- initial_prompt：初始提示词，提供上下文

Beam Search 原理：
输入音频 → Whisper 编码器 → 解码器
时间步 t=0: "你" 得分 0.8
时间步 t=1: "你好" 得分 0.75
时间步 t=2: "你好吗" 得分 0.7
最终选择得分最高的候选
""")
    
    doc.add_heading("4.3 幻觉检测", level=2)
    
    add_code_block(doc, '''REPEAT_CHAR = re.compile(r'(.)\\1{4,}')
REPEAT_WORD = re.compile(r'(\\S{2,6})\\1{3,}')
REPEAT_PHRASE = re.compile(r'(.{2,10})\\1{2,}')

def detect_prompt_echo(text: str) -> bool:
    if not text:
        return True
    if text.strip().rstrip("。") == WHISPER_INITIAL_PROMPT.rstrip("。"):
        return True
    return False

def detect_repetition(text: str) -> bool:
    if not text or len(text) < 4:
        return False
    if REPEAT_CHAR.search(text):
        return True
    if REPEAT_WORD.search(text):
        return True
    if REPEAT_PHRASE.search(text):
        return True
    return False''')
    
    add_explanation(doc, """
Whisper 常见幻觉问题：
- REPEAT_CHAR：单字符重复 5 次以上（如 "哈哈哈哈哈哈哈"）
- REPEAT_WORD：词语重复 3 次以上（如 "那个那个那个那个"）
- REPEAT_PHRASE：短语重复 2 次以上（如 "对吧对吧对吧"）
- detect_prompt_echo：检测是否只输出了初始提示词（音频为空时）
""")
    
    doc.add_heading("4.4 繁简体转换", level=2)
    
    add_code_block(doc, '''_T2S_DICT = {
    "會": "会", "來": "来", "對": "对", "開": "开",
    "學": "学", "實": "实", "現": "现", "發": "发",
}

def _traditional_to_simplified(self, text: str) -> str:
    try:
        from zhconv import convert
        return convert(text, "zh-cn")
    except ImportError:
        pass
    return ''.join(_T2S_DICT.get(char, char) for char in text)''')
    
    add_explanation(doc, """
Whisper 可能输出繁体中文（特别是台湾地区口音）。
- 优先使用 zhconv 库转换（最准确）
- 兜底使用字典映射（覆盖常用字）
""")
    
    doc.add_page_break()
    
    # ==================== 第五章 ====================
    doc.add_heading("第五章：NLP信息提取 nlp_service.py", level=1)
    
    doc.add_heading("5.1 关键词库", level=2)
    
    add_code_block(doc, '''DECISION_KEYWORDS = [
    "决定", "确定", "同意", "通过", "批准", "拍板", "敲定",
    "达成共识", "一致认为", "最终方案", "砍掉", "放弃",
]

ACTION_KEYWORDS = [
    "负责", "由.*完成", "安排.*做", "指派", "落实", "执行",
    "推进", "下一步", "上线", "启动", "搞定", "出报告",
]

ISSUE_KEYWORDS = [
    "问题", "未解决", "待讨论", "需要进一步", "遗留", "暂缓",
    "延后", "搁置", "尚需", "还需", "继续讨论", "下次会议",
]''')
    
    add_explanation(doc, """
关键词库设计：
- DECISION_KEYWORDS：匹配关键决策（如 "我们决定下周上线"）
- ACTION_KEYWORDS：匹配行动项（如 "张三负责完成开发"）
- ISSUE_KEYWORDS：匹配遗留问题（如 "服务器扩容问题待解决"）
""")
    
    doc.add_heading("5.2 正则表达式模式", level=2)
    
    add_code_block(doc, '''PERSON_PATTERNS = [
    r'(?:由|请|让|安排|指派|交给)([\\u4e00-\\u9fa5]{2,4}?)(?:负责|完成|跟进)',
    r'([\\u4e00-\\u9fa5]{2,4})(?:负责|来负责|来完成)',
    r'@([\\u4e00-\\u9fa5]{2,4})',
]

DEADLINE_PATTERNS = [
    r'(截止|deadline)[：:]\\s*(.+?)(?:[。；\\n]|$)',
    r'(\\d+月\\d+[日号])',
    r'(\\d{4}[-/]\\d{1,2}[-/]\\d{1,2})',
    r'(下周[一二三四五六日])',
]''')
    
    add_explanation(doc, """
正则表达式说明：
- PERSON_PATTERNS：提取人名
  - (?:由|请|让...)：非捕获组，匹配但不保存
  - ([\\u4e00-\\u9fa5]{2,4}?)：捕获组，保存 2-4 个中文字符
  - 示例："由张三负责" → 提取出 "张三"

- DEADLINE_PATTERNS：提取截止时间
  - 示例："截止：下周五" → 提取出 "下周五"
  - 示例："1月15日" → 提取出 "1月15日"
""")
    
    doc.add_heading("5.3 行动项提取", level=2)
    
    add_code_block(doc, '''def extract_action_items(self, text: str, _corrected: bool = False) -> list:
    sentences = self.split_sentences(text, _corrected=True)
    items = []
    seen = set()
    
    for sent in sentences:
        if not self._sentence_has_keywords(sent, _ACTION_KW_RE):
            continue

        person = ""
        for pattern in _RE_PERSON_PATTERNS:
            match = pattern.search(sent)
            if match:
                for g in match.groups():
                    if g and _RE_PERSON_CHECK.match(g):
                        person = g
                        break
                if person:
                    break

        deadline = ""
        for pattern in _RE_DEADLINE_PATTERNS:
            match = pattern.search(sent)
            if match:
                deadline = match.group(0)
                break

        cleaned = re.sub(r'^[，,。.！!？?\\s]+', '', sent)
        if not cleaned or len(cleaned) < 8:
            continue

        key = (person, cleaned[:30])
        if key in seen:
            continue
        seen.add(key)
        items.append({
            "content": cleaned,
            "responsible_person": person,
            "deadline": deadline,
            "priority": "medium"
        })
    return items''')
    
    add_explanation(doc, """
行动项提取流程：
1. 按句子分割文本
2. 检查每个句子是否包含行动项关键词
3. 提取责任人（先用正则，失败再用 jieba 词性标注）
4. 提取截止时间
5. 清理任务内容（去除标点、过滤过短内容）
6. 去重（使用 (person, content[:30]) 作为去重键）
""")
    
    doc.add_heading("5.4 参会人员提取", level=2)
    
    add_code_block(doc, '''def extract_participants(self, text: str, _corrected: bool = False) -> str:
    persons = set()
    
    # 策略 1：显式标签提取
    m = _RE_PARTICIPANT_LABEL.search(text)
    if m:
        for name in _RE_PARTICIPANT_SPLIT.split(m.group(1)):
            name = name.strip()
            if 2 <= len(name) <= 4 and _RE_PERSON_CHECK.match(name):
                persons.add(name)
    
    # 策略 2：说话人检测
    for name in self._detect_speakers(text):
        persons.add(name)
    
    # 策略 3：词性标注提取
    if self.jieba_loaded:
        for word, flag in pseg.cut(text):
            if flag == "nr" and 2 <= len(word) <= 4:
                persons.add(word)
    
    return "、".join(persons) if persons else ""''')
    
    add_explanation(doc, """
三种策略提取参会人员：
1. 显式标签：匹配 "参会人：张三、李四" 格式
2. 说话人检测：匹配 "张三说..."、"李四认为..." 格式
3. 词性标注：使用 jieba 识别词性为 nr（人名）的词
使用 set 自动去重，使用 "、" 连接返回
""")
    
    doc.add_page_break()
    
    # ==================== 第六章 ====================
    doc.add_heading("第六章：智能问答模块 qa_service.py", level=1)
    
    doc.add_heading("6.1 语义检索", level=2)
    
    add_code_block(doc, '''def _semantic_search(self, paragraphs: list, question: str, top_k: int = 3) -> list:
    if not SENTENCE_TRANSFORMERS_AVAILABLE or QAService._embedder is None:
        return paragraphs[:top_k]
    try:
        para_embeddings = QAService._embedder.encode(paragraphs, convert_to_tensor=True)
        q_embedding = QAService._embedder.encode(question, convert_to_tensor=True)
        hits = util.semantic_search(q_embedding, para_embeddings, top_k=top_k)[0]
        return [paragraphs[hit["corpus_id"]] for hit in hits]
    except Exception as e:
        return paragraphs[:top_k]''')
    
    add_explanation(doc, """
语义检索流程：
1. 使用 Sentence-BERT 将段落和问题编码为向量
2. 计算余弦相似度
3. 返回 Top-K 最相关的段落

余弦相似度公式：
cos(θ) = (A·B) / (|A|×|B|)
- 相似度越高（越接近 1），语义越相似
- 相似度越低（越接近 0），语义越不相关

降级策略：语义检索失败时返回前 N 段
""")
    
    doc.add_heading("6.2 回答生成", level=2)
    
    add_code_block(doc, '''def _generate_answer(self, question: str, context: str) -> str:
    prompt = (
        f"你是会议纪要助手。根据以下会议文本回答问题。"
        f"即使信息不完全匹配，也请基于文本内容给出最相关的回答。\\n\\n"
        f"会议文本：\\n{context}\\n\\n"
        f"问题：\\n{question}\\n\\n"
        f"请直接回答："
    )
    
    try:
        from app.services.nlp_service import local_qwen_processor
        local_qwen_processor._ensure_model()
        
        tokenizer = local_qwen_processor.__class__._tokenizer
        model = local_qwen_processor.__class__._model
        
        messages = [{"role": "user", "content": prompt}]
        input_text = tokenizer.apply_chat_template(
            messages, tokenize=False, add_generation_prompt=True, enable_thinking=False
        )
        
        import torch
        inputs = tokenizer(input_text, return_tensors="pt", truncation=True, max_length=8192)
        
        with torch.inference_mode():
            output = model.generate(
                **inputs, max_new_tokens=512, do_sample=False,
                pad_token_id=tokenizer.eos_token_id
            )
        
        response = tokenizer.decode(output[0][inputs["input_ids"].shape[1]:], skip_special_tokens=True)
        return response.strip()
    except Exception as e:
        return self._fallback_answer(question, context)''')
    
    add_explanation(doc, """
回答生成流程：
1. 构造 Prompt：角色 + 上下文 + 问题
2. 使用 Qwen3 本地模型推理
3. inputs["input_ids"].shape[1]：输入 token 数，用于切分输出
4. 降级策略：Qwen3 失败时使用关键词匹配

Prompt 设计：
- 角色设定：你是会议纪要助手
- 任务描述：根据会议文本回答问题
- 约束条件：即使信息不完全匹配，也请基于文本内容回答
""")
    
    doc.add_heading("6.3 RAG 完整流程", level=2)
    
    add_code_block(doc, '''def ask(self, meeting_id: int, question: str, db) -> dict:
    if not question or not question.strip():
        return {"success": False, "message": "问题不能为空"}

    meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
    if not meeting:
        return {"success": False, "message": f"会议 ID {meeting_id} 不存在"}

    raw_text = meeting.raw_text or ""
    if not raw_text.strip():
        return {"success": False, "message": "该会议没有原始文本，无法回答问题"}

    paragraphs = self._split_paragraphs(raw_text)
    relevant = self._semantic_search(paragraphs, question)
    context = "\\n".join(relevant)

    answer = self._generate_answer(question, context)

    return {
        "success": True,
        "answer": answer,
        "sources": relevant,
    }''')
    
    add_explanation(doc, """
RAG (Retrieval-Augmented Generation) 完整流程：
1. 验证参数
2. 获取会议原文
3. 分割段落
4. 语义检索相关段落（Retrieval）
5. 生成回答（Generation）
6. 返回结果（包含来源段落）
""")
    
    doc.add_page_break()
    
    # ==================== 第七章 ====================
    doc.add_heading("第七章：音频增强模块 audio_processor.py", level=1)
    
    doc.add_heading("7.1 降噪处理", level=2)
    
    add_code_block(doc, '''def _denoise(self, audio_segment):
    try:
        import noisereduce as nr
        samples = np.array(audio_segment.get_array_of_samples(), dtype=np.float32)
        if audio_segment.channels == 2:
            samples = samples.reshape((-1, 2)).mean(axis=1)
        reduced_samples = nr.reduce_noise(y=samples, sr=audio_segment.frame_rate)
        reduced_samples = (reduced_samples * 32767).astype(np.int16)
        return audio_segment._spawn(reduced_samples.tobytes())
    except Exception as e:
        return audio_segment''')
    
    add_explanation(doc, """
降噪处理：
- noisereduce：基于频谱门控的降噪算法
- 将 pydub 音频段转为 numpy 数组
- 立体声转单声道
- 降噪后转回 int16 格式
""")
    
    doc.add_heading("7.2 回声消除 (LMS 算法)", level=2)
    
    add_code_block(doc, '''def _echo_cancellation(self, audio_segment):
    import scipy.signal
    samples = np.array(audio_segment.get_array_of_samples(), dtype=np.float64)
    
    sr = audio_segment.frame_rate
    delay_ms = 50
    delay_samples = int(sr * delay_ms / 1000)
    filter_length = 128
    mu = 0.01

    ref = np.zeros(len(samples), dtype=np.float64)
    ref[delay_samples:] = samples[:-delay_samples]

    weights = np.zeros(filter_length, dtype=np.float64)
    output = np.zeros(len(samples), dtype=np.float64)

    for i in range(filter_length, len(samples)):
        x = ref[i - filter_length:i][::-1].copy()
        y = np.dot(weights, x)
        e = samples[i] - y
        weights += mu * e * x
        output[i] = e

    return audio_segment._spawn(output.tobytes())''')
    
    add_explanation(doc, """
LMS (Least Mean Squares) 自适应滤波器算法：

参数说明：
- delay_ms = 50：预估回声延迟（毫秒）
- filter_length = 128：滤波器长度
- mu = 0.01：步长参数（学习率）

算法流程：
1. 构造参考信号 ref（延迟版本，模拟回声）
2. 初始化滤波器权重 weights
3. 迭代更新：
   - y = np.dot(weights, x)：滤波器输出（回声估计）
   - e = samples[i] - y：误差信号（消除回声后）
   - weights += mu * e * x：更新权重
4. 输出 e 即为消除回声后的信号
""")
    
    doc.add_page_break()
    
    # ==================== 第八章 ====================
    doc.add_heading("第八章：说话人识别 speaker_diarization.py", level=1)
    
    doc.add_heading("8.1 声纹特征提取", level=2)
    
    add_code_block(doc, '''def _extract_embeddings(self, audio, segments: list) -> list:
    from resemblyzer import VoiceEncoder
    encoder = VoiceEncoder()
    
    embeddings = []
    
    for start, end in segments:
        start_sample = int(start * SAMPLE_RATE)
        end_sample = int(end * SAMPLE_RATE)
        segment = audio[start_sample:end_sample]
        
        target_length = int(1.6 * SAMPLE_RATE)
        
        if len(segment) < target_length:
            segment = np.pad(segment, (0, target_length - len(segment)))
        elif len(segment) > target_length:
            start_idx = (len(segment) - target_length) // 2
            segment = segment[start_idx:start_idx + target_length]
        
        embedding = encoder.embed_utterance(segment)
        embeddings.append(embedding)
    
    return embeddings''')
    
    add_explanation(doc, """
声纹特征提取：
- VoiceEncoder：基于 GE2E 损失训练的 LSTM 声纹编码器
- 输入：1.6 秒音频片段（16kHz 采样率）
- 输出：256 维向量
- 如果音频太短，填充静音；如果太长，截取中间部分
""")
    
    doc.add_heading("8.2 层次聚类", level=2)
    
    add_code_block(doc, '''def _cluster_speakers(self, embeddings: list) -> list:
    from sklearn.cluster import AgglomerativeClustering
    from sklearn.metrics.pairwise import cosine_similarity

    if len(embeddings) < 2:
        return [0] * len(embeddings)

    similarity_matrix = cosine_similarity(embeddings)
    distance_matrix = 1 - similarity_matrix

    n_speakers = self._estimate_speaker_count(distance_matrix)

    clustering = AgglomerativeClustering(
        n_clusters=n_speakers,
        metric='precomputed',
        linkage='average'
    )

    labels = clustering.fit_predict(distance_matrix)
    return labels''')
    
    add_explanation(doc, """
层次聚类：
- cosine_similarity：计算余弦相似度矩阵
- distance_matrix = 1 - similarity_matrix：转为距离矩阵
- AgglomerativeClustering：层次聚类
- linkage='average'：平均链接
- 自动确定说话人数量
""")
    
    doc.add_page_break()
    
    # ==================== 第九章 ====================
    doc.add_heading("第九章：实时转写模块 streaming_service.py", level=1)
    
    doc.add_heading("9.1 流式识别核心", level=2)
    
    add_code_block(doc, '''class StreamingTranscriber:
    def __init__(self, sample_rate=16000, chunk_duration=3.0):
        self.sample_rate = sample_rate
        self.chunk_duration = chunk_duration
        self.buffer = bytearray()
        self._context_text = ""

    def add_audio(self, audio_data: bytes):
        self.buffer.extend(audio_data)

    def should_transcribe(self) -> bool:
        input_rate = self._input_sample_rate or self.sample_rate
        expected_bytes = int(input_rate * 2 * (self.chunk_duration + self.overlap_duration))
        return len(self.buffer) >= expected_bytes''')
    
    add_explanation(doc, """
流式识别参数：
- sample_rate=16000：Whisper 要求 16kHz 采样率
- chunk_duration=3.0：每次识别 3 秒音频
- buffer：音频缓冲区
- _context_text：上下文文本，提升识别连贯性

缓冲区管理：
- 输入采样率: 44100 Hz (浏览器)
- 目标采样率: 16000 Hz (Whisper 要求)
- 缓冲区大小 = 44100 × 2 bytes × 3.5s ≈ 300KB
""")
    
    doc.add_heading("9.2 上下文传递", level=2)
    
    add_code_block(doc, '''def transcribe_chunk(self) -> dict:
    samples = np.frombuffer(self.buffer, dtype=np.int16)
    audio_np = samples.astype(np.float32) / 32768.0

    audio_np = self._resample_to_16k(audio_np)

    segments, _ = self._model.transcribe(
        audio_np,
        language="zh",
        beam_size=5,
        initial_prompt=self._context_text
    )

    text = "".join([seg.text for seg in segments])

    self._context_text = (self._context_text + text)[-200:]

    self.buffer = bytearray()

    return {"text": text, "is_final": False}''')
    
    add_explanation(doc, """
上下文传递机制：
- 使用上一次的识别结果作为上下文
- 提升识别连贯性（特别是长句）
- 保留最近 200 字符作为上下文
""")
    
    doc.add_page_break()
    
    # ==================== 第十章 ====================
    doc.add_heading("第十章：API路由层", level=1)
    
    doc.add_heading("10.1 REST API 设计", level=2)
    
    # 添加表格
    table = doc.add_table(rows=13, cols=4)
    table.style = "Light Grid Accent 1"
    
    headers = ["方法", "路径", "功能", "请求体"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    rows_data = [
        ["POST", "/api/meetings/upload-audio", "上传音频", "FormData"],
        ["POST", "/api/meetings/recognize", "语音识别", "Form"],
        ["POST", "/api/meetings", "创建会议", "JSON"],
        ["GET", "/api/meetings", "会议列表", "Query"],
        ["GET", "/api/meetings/{id}", "会议详情", "-"],
        ["PUT", "/api/meetings/{id}", "更新会议", "JSON"],
        ["DELETE", "/api/meetings/{id}", "删除会议", "-"],
        ["POST", "/api/meetings/{id}/generate-minutes", "生成纪要", "-"],
        ["POST", "/api/meetings/{id}/qa", "智能问答", "JSON"],
        ["POST", "/api/meetings/{id}/qa/stream", "流式问答", "JSON"],
        ["GET", "/api/export/docx/{id}", "导出Word", "-"],
        ["GET", "/api/export/markdown/{id}", "导出Markdown", "-"],
    ]
    
    for i, row_data in enumerate(rows_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("10.2 依赖注入示例", level=2)
    
    add_code_block(doc, '''@router.get("")
async def list_meetings(
    page: int = 1,
    size: int = 20,
    db: Session = Depends(get_db),
    user: User = Depends(require_user),
):
    meetings = db.query(Meeting).filter(Meeting.created_by == user.id).all()
    return APIResponse(success=True, data={"items": meetings})''')
    
    add_explanation(doc, """
FastAPI 依赖注入：
- page/size：Query 参数
- db: Session = Depends(get_db)：数据库会话（自动注入）
- user: User = Depends(require_user)：当前用户（自动认证）
""")
    
    doc.add_page_break()
    
    # ==================== 第十一章 ====================
    doc.add_heading("第十一章：前端交互逻辑 app.js", level=1)
    
    doc.add_heading("11.1 核心类设计", level=2)
    
    add_code_block(doc, '''class MeetingApp {
    constructor() {
        this.currentMeetingId = null;
        this.currentMinutes = null;
        this.actionItems = [];
        this.currentUser = null;
        this.sessionToken = null;
        this.isRecording = false;
        this._mediaRecorder = null;
        this._audioChunks = [];
        this._ws = null;
    }
    
    async init() {
        this._bindEvents();
        this._bindKeyboardShortcuts();
        this._loadDraft();
        this._setupAutoSave();
        await this._checkLoginStatus();
        await this.loadHistory();
    }
}''', language="javascript")
    
    add_explanation(doc, """
前端核心类：
- 状态管理：当前会议、纪要、行动项、用户信息
- 录音相关：MediaRecorder、音频块
- WebSocket：实时转写连接
- init()：初始化方法，绑定事件、加载草稿、检查登录状态
""")
    
    doc.add_heading("11.2 事件绑定", level=2)
    
    add_code_block(doc, '''_bindEvents() {
    const dropZone = document.getElementById('dropZone');
    const fileInput = document.getElementById('fileInput');

    dropZone.addEventListener('click', () => fileInput.click());
    dropZone.addEventListener('dragover', (e) => {
        e.preventDefault();
        dropZone.classList.add('drag-over');
    });
    dropZone.addEventListener('drop', (e) => {
        e.preventDefault();
        if (e.dataTransfer.files.length) this.handleFile(e.dataTransfer.files[0]);
    });

    document.getElementById('btnSubmitText').addEventListener('click', () => this.submitText());
    document.getElementById('btnGenerateMinutes').addEventListener('click', () => this.generateMinutes());
}''', language="javascript")
    
    add_explanation(doc, """
事件绑定：
- 拖拽上传：dragover、dragleave、drop 事件
- 点击上传：click 事件
- 按钮点击：submitText、generateMinutes 等
""")
    
    doc.add_heading("11.3 进度条管理", level=2)
    
    add_code_block(doc, '''showProgress(label, percent, detail) {
    this._els.progressSection.style.display = 'block';
    this._els.progressLabel.textContent = label;
    this._els.progressPercent.textContent = percent + '%';
    this._els.progressBar.style.width = percent + '%';
    this._els.progressDetail.textContent = detail;
}

_startProgressAnimation() {
    const stages = [
        { delay: 3000, pct: 20, detail: '正在处理音频数据...' },
        { delay: 8000, pct: 35, detail: '正在分析语音模式...' },
        { delay: 15000, pct: 50, detail: '正在转换为文字...' },
        { delay: 25000, pct: 65, detail: '正在优化转写结果...' },
    ];
}''', language="javascript")
    
    add_explanation(doc, """
进度条管理：
- showProgress：显示进度条
- updateProgress：更新进度
- _startProgressAnimation：动态进度更新
""")

    doc.add_page_break()

    # ==================== 第十二章：SSE 流式输出 ====================
    doc.add_heading("第十二章：SSE 流式输出 qa.py", level=1)

    doc.add_heading("12.1 SSE 协议原理", level=2)

    add_explanation(doc, """
SSE (Server-Sent Events) 协议说明：

1. 定义：
   - 服务端单向推送数据给客户端
   - 基于 HTTP 协议，不需要 WebSocket
   - Content-Type: text/event-stream

2. 数据格式：
   data: {json}\n\n

3. 每条消息结构：
   - 以 "data: " 开头
   - 后接 JSON 数据
   - 以两个换行符 (\n\n) 结尾

4. 优势：
   - 用户可以实时看到生成过程，提升体验
   - 不需要等待完整响应，首字延迟低
   - 可以实现打字机效果
   - 比 WebSocket 更简单，不需要额外协议

5. 适用场景：
   - AI 文本生成（如 ChatGPT 流式输出）
   - 实时日志推送
   - 股票行情推送
   - 新闻滚动更新
""")

    doc.add_heading("12.2 流式问答接口代码", level=2)

    add_code_block(doc, '''@qa_router.post("/{meeting_id}/qa/stream")
async def ask_question_stream(
    meeting_id: int,
    request: QARequest,
    db: Session = Depends(get_db),
    user: User = Depends(require_user),
):
    """流式问答接口 - 使用 SSE 协议"""
    
    def generate():
        """
        生成器函数 - 用于 SSE 流式输出
        
        SSE 数据格式：
        data: {"text": "根据"}\n\n
        data: {"text": "会议"}\n\n
        data: {"text": "记录"}\n\n
        data: {"done": true}\n\n
        """
        try:
            # 1. 验证会议是否存在
            from app.models import Meeting
            meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
            if not meeting:
                yield f"data: {json.dumps({'error': '会议不存在'})}\\n\\n"
                return

            # 2. 获取会议原文
            raw_text = meeting.raw_text or ""
            if not raw_text.strip():
                yield f"data: {json.dumps({'error': '该会议没有原始文本'})}\\n\\n"
                return

            # 3. 语义检索相关段落
            paragraphs = qa_service._split_paragraphs(raw_text)
            relevant = qa_service._semantic_search(paragraphs, request.question)
            context = "\\n".join(relevant)

            # 4. 构造 Prompt
            prompt = (
                f"你是会议纪要助手。根据以下会议文本回答问题。\\n\\n"
                f"会议文本：\\n{context}\\n\\n"
                f"问题：{request.question}\\n\\n"
                f"请直接回答："
            )

            # 5. 加载 Qwen3 模型
            from app.services.nlp_service import local_qwen_processor
            local_qwen_processor._ensure_model()
            
            tokenizer = local_qwen_processor.__class__._tokenizer
            model = local_qwen_processor.__class__._model

            # 6. 使用 Qwen3 流式生成
            messages = [{"role": "user", "content": prompt}]
            input_text = tokenizer.apply_chat_template(
                messages, tokenize=False, add_generation_prompt=True,
                enable_thinking=False
            )

            import torch
            from transformers import TextIteratorStreamer
            import threading

            inputs = tokenizer(input_text, return_tensors="pt", 
                              truncation=True, max_length=8192)
            
            # 创建流式输出器
            streamer = TextIteratorStreamer(
                tokenizer, skip_prompt=True, skip_special_tokens=True
            )

            gen_kwargs = dict(
                **inputs,
                max_new_tokens=512,
                do_sample=False,
                pad_token_id=tokenizer.eos_token_id,
                streamer=streamer,
            )

            # 在后台线程运行推理
            def _generate():
                with torch.inference_mode():
                    model.generate(**gen_kwargs)

            gen_thread = threading.Thread(target=_generate, daemon=True)
            gen_thread.start()

            # 7. 逐 token 推送给客户端
            for new_text in streamer:
                if new_text.strip():
                    yield f"data: {json.dumps({'text': new_text})}\\n\\n"

            # 发送完成信号
            yield f"data: {json.dumps({'done': True})}\\n\\n"

        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\\n\\n"

    return StreamingResponse(generate(), media_type="text/event-stream")''')

    add_explanation(doc, """
代码详解：

1. @qa_router.post("/{meeting_id}/qa/stream")
   - 定义 POST 请求的流式问答端点
   - 返回 StreamingResponse，Content-Type 为 text/event-stream

2. generate() 生成器函数
   - 使用 yield 关键字逐块输出数据
   - 每次 yield 发送一条 SSE 消息

3. TextIteratorStreamer
   - HuggingFace 提供的流式输出器
   - 将模型生成的 token 逐个输出
   - skip_prompt=True：跳过 prompt 部分
   - skip_special_tokens=True：跳过特殊 token

4. 后台线程推理
   - 使用 threading.Thread 在后台运行模型推理
   - 主线程从 streamer 读取输出并推送给客户端
   - 这样可以实现真正的流式输出

5. SSE 数据格式
   - data: {"text": "根据"}\n\n
   - data: {"text": "会议"}\n\n
   - data: {"done": true}\n\n
   - 客户端通过 EventSource API 接收
""")

    doc.add_heading("12.3 SSE 通信流程图", level=2)

    add_explanation(doc, """
SSE 流式问答通信流程：

客户端 (浏览器)                    服务端 (FastAPI)
     │                                  │
     │  POST /qa/stream                 │
     │  {"question": "..."}             │
     │  ──────────────────────────────> │
     │                                  │
     │  HTTP 200                        │
     │  Content-Type: text/event-stream │
     │  <────────────────────────────── │
     │                                  │
     │  data: {"text": "根据"}           │
     │  <────────────────────────────── │
     │                                  │
     │  data: {"text": "会议"}           │
     │  <────────────────────────────── │
     │                                  │
     │  data: {"text": "记录"}           │
     │  <────────────────────────────── │
     │                                  │
     │  data: {"text": "分析"}           │
     │  <────────────────────────────── │
     │                                  │
     │  data: {"done": true}            │
     │  <────────────────────────────── │
     │                                  │
     │  连接关闭                         │
     │  <────────────────────────────── │
""")

    doc.add_heading("12.4 前端 SSE 接收代码", level=2)

    add_code_block(doc, '''// 前端使用 EventSource API 接收 SSE
async askQuestion() {
    const question = document.getElementById('qaInput').value;
    
    // 使用 fetch 发送请求
    const response = await fetch(`/api/meetings/${this.currentMeetingId}/qa/stream`, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ question }),
    });
    
    // 创建读取器
    const reader = response.body.getReader();
    const decoder = new TextDecoder();
    
    while (true) {
        const { done, value } = await reader.read();
        if (done) break;
        
        // 解析 SSE 数据
        const text = decoder.decode(value);
        const lines = text.split('\\n');
        
        for (const line of lines) {
            if (line.startsWith('data: ')) {
                const data = JSON.parse(line.slice(6));
                
                if (data.text) {
                    // 追加文本到回答区域（打字机效果）
                    this.appendAnswer(data.text);
                } else if (data.done) {
                    // 流式输出完成
                    console.log('流式输出完成');
                } else if (data.error) {
                    // 发生错误
                    console.error('错误:', data.error);
                }
            }
        }
    }
}''', language="javascript")

    add_explanation(doc, """
前端 SSE 接收流程：

1. 使用 fetch 发送 POST 请求
2. 获取 Response.body 的 ReadableStream
3. 创建 TextDecoder 解码二进制数据
4. 循环读取数据块
5. 解析 SSE 格式（data: {json}）
6. 根据数据类型处理：
   - text：追加文本到回答区域（打字机效果）
   - done：流式输出完成
   - error：显示错误信息

打字机效果实现：
- 每收到一个 text 数据块
- 追加到回答区域的末尾
- 自动滚动到底部
""")

    doc.add_heading("12.5 WebSocket vs SSE 对比", level=2)

    # 添加对比表格
    table = doc.add_table(rows=8, cols=3)
    table.style = "Light Grid Accent 1"

    headers = ["特性", "WebSocket", "SSE"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    rows_data = [
        ["协议", "ws:// 或 wss://", "http:// 或 https://"],
        ["方向", "双向通信", "服务端单向推送"],
        ["数据格式", "任意格式", "文本 (text/event-stream)"],
        ["复杂度", "较高（需要握手）", "较低（基于 HTTP）"],
        ["浏览器支持", "所有现代浏览器", "除 IE 外所有现代浏览器"],
        ["断线重连", "需要手动实现", "浏览器自动重连"],
        ["适用场景", "聊天、游戏、实时协作", "AI 生成、日志推送、行情"],
    ]

    for i, row_data in enumerate(rows_data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data

    doc.add_paragraph()

    add_explanation(doc, """
WebSocket vs SSE 选择建议：

使用 SSE 的场景：
- 服务端单向推送（如 AI 文本生成）
- 需要自动断线重连
- 简单的实时更新

使用 WebSocket 的场景：
- 双向通信（如聊天应用）
- 需要低延迟
- 复杂的实时协作

本项目选择 SSE 的原因：
1. 问答是单向的（用户问，AI 答）
2. SSE 自动断线重连，更稳定
3. 实现更简单，不需要额外协议
""")

    doc.add_heading("12.6 TextIteratorStreamer 工作原理", level=2)

    add_explanation(doc, """
TextIteratorStreamer 是 HuggingFace Transformers 库提供的流式输出器。

工作原理：
┌─────────────────────────────────────────────────────────────────┐
│                    TextIteratorStreamer 工作流程                  │
├─────────────────────────────────────────────────────────────────┤
│                                                                 │
│  后台线程（模型推理）                                            │
│       │                                                         │
│       ▼                                                         │
│  ┌─────────────────┐                                           │
│  │  model.generate  │                                           │
│  │  (生成 token)    │                                           │
│  └────────┬────────┘                                           │
│           │                                                     │
│           ▼                                                     │
│  ┌─────────────────┐                                           │
│  │  streamer.put()  │  ← 将 token 放入队列                     │
│  └────────┬────────┘                                           │
│           │                                                     │
│           ▼                                                     │
│  ┌─────────────────┐                                           │
│  │   token 队列     │  ← 线程安全的队列                         │
│  └────────┬────────┘                                           │
│           │                                                     │
│           │  主线程（SSE 输出）                                  │
│           │       │                                             │
│           ▼       ▼                                             │
│  ┌─────────────────┐                                           │
│  │  streamer.get()  │  ← 从队列读取 token                      │
│  └────────┬────────┘                                           │
│           │                                                     │
│           ▼                                                     │
│  ┌─────────────────┐                                           │
│  │  yield SSE 消息  │  ← 推送给客户端                           │
│  └─────────────────┘                                           │
│                                                                 │
└─────────────────────────────────────────────────────────────────┘

关键点：
1. 后台线程负责模型推理，生成 token
2. 主线程负责读取 token 并推送给客户端
3. 通过线程安全的队列进行通信
4. 实现了真正的流式输出
""")

    doc.add_heading("12.7 后台线程的作用", level=2)

    add_code_block(doc, '''# 为什么要使用后台线程？

# 错误示例（不使用后台线程）：
# def generate():
#     output = model.generate(**inputs)  # 阻塞！主线程卡住
#     for token in output:
#         yield f"data: {token}\\n\\n"

# 正确示例（使用后台线程）：
def generate():
    # 1. 创建流式输出器
    streamer = TextIteratorStreamer(tokenizer)
    
    # 2. 定义推理函数
    def _generate():
        with torch.inference_mode():
            model.generate(**inputs, streamer=streamer)
    
    # 3. 在后台线程运行推理
    gen_thread = threading.Thread(target=_generate)
    gen_thread.start()
    
    # 4. 主线程读取 streamer（不阻塞）
    for new_text in streamer:
        yield f"data: {json.dumps({'text': new_text})}\\n\\n"''')

    add_explanation(doc, """
为什么需要后台线程？

问题：
- model.generate() 是阻塞调用
- 如果在主线程调用，主线程会被阻塞
- 无法同时读取 streamer 的输出

解决方案：
- 在后台线程运行 model.generate()
- 主线程从 streamer 读取输出
- 通过线程安全的队列通信

这样可以实现：
- 后台线程：模型推理，生成 token
- 主线程：读取 token，推送 SSE 消息
- 两者并行，实现真正的流式输出
""")

    doc.add_heading("12.8 SSE 消息格式详解", level=2)

    add_explanation(doc, """
SSE 消息格式：

1. 文本消息：
   data: {"text": "根据"}\n\n

2. 完成信号：
   data: {"done": true}\n\n

3. 错误消息：
   data: {"error": "会议不存在"}\n\n

消息结构：
┌─────────────────────────────────────────┐
│  data: {JSON数据}\n\n                    │
│  ↑     ↑            ↑  ↑               │
│  │     │            │  └── 两个换行符   │
│  │     │            └── JSON 数据       │
│  │     └── 固定前缀                     │
│  └── 固定前缀                           │
└─────────────────────────────────────────┘

为什么用两个换行符？
- SSE 协议规定：每条消息以两个换行符结尾
- 浏览器通过两个换行符判断消息边界
""")

    doc.add_heading("12.9 完整的 SSE 通信流程", level=2)

    add_explanation(doc, """
完整的 SSE 流式问答流程：

时间线：
────────────────────────────────────────────────────────────────►

客户端                    服务端
  │                         │
  │  1. POST /qa/stream     │
  │  {"question": "..."}    │
  │  ─────────────────────>│
  │                         │
  │                         │  2. 验证会议
  │                         │  3. 语义检索
  │                         │  4. 构造 Prompt
  │                         │  5. 加载模型
  │                         │
  │  6. HTTP 200            │
  │  Content-Type:         │
  │  text/event-stream     │
  │  <─────────────────────│
  │                         │
  │  7. data: {"text": "根"}│
  │  <─────────────────────│
  │  显示 "根"              │
  │                         │
  │  8. data: {"text": "据"}│
  │  <─────────────────────│
  │  显示 "根据"            │
  │                         │
  │  9. data: {"text": "会"}│
  │  <─────────────────────│
  │  显示 "根据会"          │
  │                         │
  │  10. data: {"text": "议"}│
  │  <─────────────────────│
  │  显示 "根据会议"        │
  │                         │
  │  ... 继续推送 ...       │
  │                         │
  │  data: {"done": true}   │
  │  <─────────────────────│
  │  流式输出完成            │
  │                         │
  │  11. 连接关闭           │
  │  <─────────────────────│
""")

    # 保存文档
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "代码讲解文档_最终版.docx")
    doc.save(output_path)
    print(f"Word 文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    generate_docx()
