# -*- coding: utf-8 -*-
"""会议纪要生成与导出服务"""
import os
import re
from datetime import datetime


class MinutesGenerator:
    """纪要生成器"""

    TEMPLATE = """
# {title}

**会议时间**：{date}
**参会人员**：{participants}

---

## 一、会议议题
{topic}

## 二、会议概述
{overview}

## 三、讨论要点
{discussion_points}

## 四、关键决策
{key_decisions}

## 五、行动项
{action_items}

## 六、遗留问题
{unresolved_issues}

## 七、会议总结
{summary}

---
*纪要由会议纪要智能体自动生成*
*生成时间：{generated_at}*
"""

    def generate_markdown(self, meeting_data: dict) -> str:
        action_items = meeting_data.get("action_items", [])
        if action_items:
            action_lines = []
            for item in action_items:
                line = f"- [ ] {item['content']}"
                if item.get('responsible_person'):
                    line += f" —— **负责人**：{item['responsible_person']}"
                if item.get('deadline'):
                    line += f" **截止时间**：{item['deadline']}"
                if item.get('priority'):
                    line += f" [{item['priority']}]"
                action_lines.append(line)
            action_text = "\n".join(action_lines)
        else:
            action_text = "暂无行动项"

        meeting_date = meeting_data.get("meeting_date", "")
        if not meeting_date:
            meeting_date = datetime.now().strftime("%Y年%m月%d日")

        # 检查是否有模板结构
        template_structure = meeting_data.get("template_structure")
        if template_structure:
            # 使用自定义模板结构生成
            return self._generate_with_template(meeting_data, template_structure, action_text)

        return self.TEMPLATE.format(
            title=meeting_data.get("title", "会议纪要"),
            date=meeting_date,
            participants=meeting_data.get("participants", "待补充"),
            topic=meeting_data.get("topic", ""),
            overview=meeting_data.get("overview", ""),
            discussion_points=meeting_data.get("discussion_points", ""),
            key_decisions=meeting_data.get("key_decisions", ""),
            action_items=action_text,
            unresolved_issues=meeting_data.get("unresolved_issues", ""),
            summary=meeting_data.get("summary", ""),
            generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        )

    def _generate_with_template(self, meeting_data: dict, template_structure: str, action_text: str) -> str:
        """使用自定义模板结构生成纪要"""
        meeting_date = meeting_data.get("meeting_date", "")
        if not meeting_date:
            meeting_date = datetime.now().strftime("%Y年%m月%d日")

        # 解析模板结构，提取各部分标题
        sections = []
        for line in template_structure.split("\n"):
            line = line.strip()
            if line:
                # 移除序号前缀
                import re
                clean_line = re.sub(r'^[\d.、]+\s*', '', line)
                sections.append(clean_line)

        # 根据模板结构生成内容
        content_parts = []
        content_parts.append(f"# {meeting_data.get('title', '会议纪要')}")
        content_parts.append("")
        content_parts.append(f"**会议时间**：{meeting_date}")
        content_parts.append(f"**参会人员**：{meeting_data.get('participants', '待补充')}")
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")

        # 映射模板章节到内容
        content_mapping = {
            "工作总结": meeting_data.get("overview", ""),
            "工作汇报": meeting_data.get("overview", ""),
            "问题": meeting_data.get("unresolved_issues", ""),
            "解决方案": meeting_data.get("key_decisions", ""),
            "工作计划": meeting_data.get("summary", ""),
            "计划": meeting_data.get("summary", ""),
            "协调": meeting_data.get("discussion_points", ""),
            "进度": meeting_data.get("overview", ""),
            "风险": meeting_data.get("unresolved_issues", ""),
            "障碍": meeting_data.get("unresolved_issues", ""),
            "资源": meeting_data.get("discussion_points", ""),
            "里程碑": meeting_data.get("key_decisions", ""),
            "需求": meeting_data.get("topic", ""),
            "背景": meeting_data.get("overview", ""),
            "功能": meeting_data.get("discussion_points", ""),
            "技术": meeting_data.get("discussion_points", ""),
            "排期": meeting_data.get("summary", ""),
            "确认": meeting_data.get("key_decisions", ""),
            "主题": meeting_data.get("topic", ""),
            "内容": meeting_data.get("discussion_points", ""),
            "案例": meeting_data.get("discussion_points", ""),
            "讨论": meeting_data.get("discussion_points", ""),
            "学习": meeting_data.get("summary", ""),
            "议题": meeting_data.get("topic", ""),
            "创意": meeting_data.get("discussion_points", ""),
            "方案": meeting_data.get("key_decisions", ""),
            "行动": action_text,
            "跟进": meeting_data.get("summary", ""),
            "决策": meeting_data.get("key_decisions", ""),
            "要点": meeting_data.get("discussion_points", ""),
            "总结": meeting_data.get("summary", ""),
        }

        for section in sections:
            content_parts.append(f"## {section}")
            content_parts.append("")
            # 查找匹配的内容
            matched_content = ""
            for keyword, content in content_mapping.items():
                if keyword in section:
                    matched_content = content
                    break
            if matched_content:
                content_parts.append(matched_content)
            else:
                content_parts.append("（待补充）")
            content_parts.append("")

        # 添加行动项
        content_parts.append("## 行动项")
        content_parts.append("")
        content_parts.append(action_text)
        content_parts.append("")
        content_parts.append("---")
        content_parts.append(f"*纪要由会议纪要智能体自动生成*")
        content_parts.append(f"*生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")

        return "\n".join(content_parts)

    def generate_html(self, meeting_data: dict) -> str:
        md = self.generate_markdown(meeting_data)
        html = md
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        html = re.sub(r'- \[ \] (.+)', r'<li><input type="checkbox"> \1</li>', html)
        html = re.sub(r'^• (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
        html = re.sub(r'^---$', r'<hr>', html, flags=re.MULTILINE)
        html = html.replace('\n', '<br>\n')
        return f'<div class="meeting-minutes">{html}</div>'

    def export_to_docx(self, meeting_data: dict, output_path: str) -> str:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn, nsdecls
            from docx.oxml import parse_xml

            doc = Document()

            # ============ 页面设置 ============
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

            # ============ 全局样式（公文字体） ============
            style = doc.styles['Normal']
            style.font.name = '仿宋'
            style.font.size = Pt(16)  # 三号字
            style.paragraph_format.line_spacing = Pt(28)  # 固定值28磅
            style.paragraph_format.space_after = Pt(4)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

            # Heading 1 - 方正小标宋简体 二号
            if 'Heading 1' in doc.styles:
                h1 = doc.styles['Heading 1']
                h1.font.name = '方正小标宋简体'
                h1.font.size = Pt(22)  # 二号字
                h1.font.bold = False
                h1.font.color.rgb = RGBColor(0, 0, 0)
                h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                h1.paragraph_format.space_before = Pt(0)
                h1.paragraph_format.space_after = Pt(28)
                h1.paragraph_format.line_spacing = Pt(32)
                h1.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

            # Heading 2 - 黑体 三号
            if 'Heading 2' in doc.styles:
                h2 = doc.styles['Heading 2']
                h2.font.name = '黑体'
                h2.font.size = Pt(16)  # 三号字
                h2.font.bold = False
                h2.font.color.rgb = RGBColor(0, 0, 0)
                h2.paragraph_format.space_before = Pt(28)
                h2.paragraph_format.space_after = Pt(0)
                h2.paragraph_format.line_spacing = Pt(28)
                h2.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

            # Heading 3 - 楷体 三号
            if 'Heading 3' in doc.styles:
                h3 = doc.styles['Heading 3']
                h3.font.name = '楷体'
                h3.font.size = Pt(16)  # 三号字
                h3.font.bold = False
                h3.font.color.rgb = RGBColor(0, 0, 0)
                h3.paragraph_format.space_before = Pt(14)
                h3.paragraph_format.space_after = Pt(0)
                h3.paragraph_format.line_spacing = Pt(28)
                h3.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

            # ============ Helper functions ============
            def _set_shading(cell, color_hex):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            def _add_p(text, font_name=None, size=Pt(16), bold=False, color=None,
                       alignment=None, space_before=0, space_after=4,
                       line_spacing=Pt(28), first_indent=None):
                if font_name is None:
                    font_name = '仿宋'
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(space_before)
                p.paragraph_format.space_after = Pt(space_after)
                p.paragraph_format.line_spacing = line_spacing
                if alignment is not None:
                    p.alignment = alignment
                if first_indent is not None:
                    p.paragraph_format.first_line_indent = Cm(first_indent)
                run = p.add_run(text)
                run.font.name = font_name
                run.font.size = size
                run.font.bold = bold
                if color:
                    run.font.color.rgb = color
                run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                return p

            def _add_bullets(text_block, font_name=None):
                if font_name is None:
                    font_name = '仿宋'
                skip_values = {'无', '无决策', '无遗留问题',
                    '本次会议未做出明确决策',
                    '详见会议原始记录',
                    '会议中未明确记录关键决策'}
                if not text_block or text_block.strip() in skip_values:
                    _add_p(text_block)
                    return
                for line in text_block.strip().split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    clean = re.sub(r'^[\u2022\-*]\s*', '', line).strip()
                    if clean:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.line_spacing = Pt(28)
                        run = p.add_run(clean)
                        run.font.name = font_name
                        run.font.size = Pt(16)
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

            def _add_hr():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(8)
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
                pPr.append(pBdr)

            # ============ 1. Main title ============
            title_text = meeting_data.get('title', '会议纪要')
            doc.add_heading(title_text, level=1)

            meeting_date = meeting_data.get('meeting_date', '')
            participants = meeting_data.get('participants', '待补充')
            subtitle = f'会议时间：{meeting_date}    参会人员：{participants}'
            _add_p(subtitle, size=Pt(14), color=RGBColor(0x66, 0x66, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)
            _add_hr()

            # ============ 2. Content sections ============
            sections_data = [
                ('一、会议议题', meeting_data.get('topic', ''), False),
                ('二、会议概述', meeting_data.get('overview', ''), False),
                ('三、讨论要点', meeting_data.get('discussion_points', ''), True),
                ('四、关键决策', meeting_data.get('key_decisions', ''), True),
                ('五、行动项', None, True),
                ('六、遗留问题', meeting_data.get('unresolved_issues', ''), True),
                ('七、会议总结', meeting_data.get('summary', ''), False),
            ]

            for sec_title, sec_content, use_bullets in sections_data:
                doc.add_heading(sec_title, level=2)

                if sec_title == '五、行动项':
                    action_items = meeting_data.get('action_items', [])
                    if action_items:
                        table = doc.add_table(rows=1 + len(action_items), cols=4)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.style = 'Table Grid'

                        for row in table.rows:
                            row.cells[0].width = Cm(7.5)
                            row.cells[1].width = Cm(2.5)
                            row.cells[2].width = Cm(2.5)
                            row.cells[3].width = Cm(1.5)

                        headers = ['行动内容', '负责人', '截止时间', '优先级']
                        for i, h_text in enumerate(headers):
                            cell = table.rows[0].cells[i]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.paragraph_format.line_spacing = Pt(28)
                            run = p.add_run(h_text)
                            run.font.name = '黑体'
                            run.font.size = Pt(14)
                            run.font.bold = False
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            _set_shading(cell, '1A3C6E')

                        for row_idx, item in enumerate(action_items):
                            row = table.rows[row_idx + 1]
                            vals = [item.get('content', ''), item.get('responsible_person', ''), item.get('deadline', ''), item.get('priority', '')]
                            for col_idx, val in enumerate(vals):
                                cell = row.cells[col_idx]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                                p.paragraph_format.line_spacing = Pt(28)
                                run = p.add_run(val)
                                run.font.name = '仿宋'
                                run.font.size = Pt(14)
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                            if row_idx % 2 == 1:
                                for col_idx in range(4):
                                    _set_shading(row.cells[col_idx], 'F2F6FC')
                        doc.add_paragraph()
                    else:
                        _add_p('暂无行动项', color=RGBColor(0x99, 0x99, 0x99))

                elif use_bullets and sec_content:
                    _add_bullets(sec_content)
                elif sec_content:
                    _add_p(sec_content, first_indent=0.74)

            # ============ 3. Footer ============
            doc.add_paragraph()
            _add_hr()
            _add_p(f'纪要由会议纪要智能体自动生成 | {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', size=Pt(10), color=RGBColor(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0)

            doc.save(output_path)
            return output_path
        except ImportError:
            raise ImportError('导出Word需要安装python-docx: pip install python-docx')


minutes_generator = MinutesGenerator()
